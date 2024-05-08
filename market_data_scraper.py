"""
Market Data Scraper

This script is designed to scrape various types of market data from different sources such as Yahoo Finance, MarketBeat, and others. It includes functionalities to retrieve data on stock prices, futures, cryptocurrencies, indices, and unusual options volumes, among others. The data can be used for market analysis, investment decision-making, or financial reporting.

The script is structured into classes and functions, each handling a specific type of data or a specific source. It utilizes libraries such as requests, BeautifulSoup, pandas, and yfinance to fetch and process the data.
"""

import yfinance
import pandas as pd
from bs4 import BeautifulSoup
import requests
import time
import datetime
import logging
import os
import traceback
import json  # Added for JSON support

# Initialize logging
logging.basicConfig(filename='market_data_scraper.log', level=logging.INFO, format='%(asctime)s:%(levelname)s:%(message)s')

# Market Data Scraper Class
# This class contains methods to scrape market data from various sources.
class MarketDataScraper:

    def __init__(self):
        self._session = requests.Session()

        self._marketbeat_unusual_calls_vol_url = \
            'https://www.marketbeat.com/market-data/unusual-call-options-volume/'

        # This url is for unusual Put options activity
        self._marketbeat_unusual_puts_vol_url = \
            'https://www.marketbeat.com/market-data/unusual-put-options-volume/'

        # The url is for index futures, and commodity futures
        self._yf_futures_url = 'https://finance.yahoo.com/commodities'

        # This url is for the Index data table
        self._yf_index_data_url = 'https://finance.yahoo.com/world-indices'

        # This url is for Crypto Data table
        self._yf_crypto_data_url = 'https://finance.yahoo.com/cryptocurrencies'

        # This url is for Trending Tickers table (most searched)
        self._yf_trending_tickers_url = 'https://finance.yahoo.com/trending-tickers'

        # This url is for the Most Active tickers
        self._yf_most_active_url = 'https://finance.yahoo.com/most-active'

        # This url is for the top Gaining tickers
        self._yf_top_gainers_url = 'https://finance.yahoo.com/gainers'

        # This url is for the top Losing tickers
        self._yf_top_losers_url = 'https://finance.yahoo.com/losers'

        # This url is for the major currencies data table on yfinance
        self._yf_currencies_url = 'https://finance.yahoo.com/currencies'

        # url for sp 500 tickers sorted by index weight
        self._slick_charts_sp_500_url = 'https://www.slickcharts.com/sp500'

    def make_soup(self, url):
        data = False
        try:
            raw_page = self._session.get(url)
            data = BeautifulSoup(raw_page.content, 'lxml')
        except Exception as e:
            logging.exception(f'Error making soup for URL: {url}', exc_info=True)
        finally:
            return data

    # Method to export data to file
    def export_data_to_file(self, data, file_format, file_name):
        """
        Exports data to a file in either HTML, docx, or JSON format.
        :param data: Data to be exported.
        :param file_format: Format of the file ('html', 'docx', or 'json').
        :param file_name: Name of the file to be created.
        """
        try:
            if file_format == 'html':
                # Export data to HTML using pandas
                data.to_html(f'{file_name}.html')
                logging.info(f'Data exported to {file_name}.html successfully.')
            elif file_format == 'docx':
                # Export data to docx using python-docx
                from docx import Document
                doc = Document()
                doc.add_heading('Data Export', 0)
                table = doc.add_table(rows=1, cols=len(data.columns))
                hdr_cells = table.rows[0].cells
                for i, column in enumerate(data.columns):
                    hdr_cells[i].text = str(column)
                for index, row in data.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
                doc.save(f'{file_name}.docx')
                logging.info(f'Data exported to {file_name}.docx successfully.')
            elif file_format == 'json':
                # Export data to JSON
                data.to_json(f'{file_name}.json', orient='records', lines=True)
                logging.info(f'Data exported to {file_name}.json successfully.')
            else:
                logging.error('Unsupported file format. Please choose either "html", "docx", or "json".')
        except Exception as e:
            logging.exception('Failed to export data', exc_info=True)

if __name__ == '__main__':
    pass
